from docx import Document
from io import BytesIO

class DOCXReportGenerator:
    def __init__(self):
        pass

    def generate(self, context: dict) -> bytes:
        doc = Document()
        doc.add_heading(context.get('title', 'Reporte'), level=1)
        if context.get('summary'):
            doc.add_paragraph(context['summary'])
        doc.add_paragraph(f"Fecha desde: {context.get('date_range_start')} hasta: {context.get('date_range_end')}")
        doc.add_paragraph(f"Unidad de negocio: {context.get('business_unit')}")
        doc.add_paragraph(f"País: {context.get('country')}")
        doc.add_paragraph(f"Idioma: {context.get('language')}")

        # Add articles as bullet points
        articles = context.get('articles', [])
        if articles:
            doc.add_heading('Artículos incluidos', level=2)
            for article_id in articles:
                doc.add_paragraph(str(article_id), style='ListBullet')

        output = BytesIO()
        doc.save(output)
        return output.getvalue()
